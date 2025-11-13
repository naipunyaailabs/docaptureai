from fastapi import FastAPI, File, UploadFile
from fastapi.responses import JSONResponse
import base64
import tempfile
import os
from docx import Document
from typing import Dict, Any, List

app = FastAPI()

def decode_base64_file(base64_data: str) -> bytes:
    """Decode base64 encoded data to bytes"""
    # Remove data URL prefix if present
    if base64_data.startswith('data:'):
        base64_data = base64_data.split(',', 1)[1]
    
    # Decode base64 data
    return base64.b64decode(base64_data)

def extract_word_content(file_bytes: bytes) -> Dict[str, Any]:
    """Extract content from Word document bytes"""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
        tmp_file.write(file_bytes)
        tmp_file_path = tmp_file.name
    
    try:
        # Open the document
        doc = Document(tmp_file_path)
        
        # Extract text content
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                paragraphs.append(paragraph.text)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables.append(table_data)
        
        # Extract metadata
        core_properties = doc.core_properties
        metadata = {
            'title': core_properties.title,
            'author': core_properties.author,
            'created': core_properties.created,
            'modified': core_properties.modified,
            'subject': core_properties.subject,
            'keywords': core_properties.keywords
        }
        
        return {
            'text_content': '\n'.join(paragraphs),
            'paragraph_count': len(paragraphs),
            'table_count': len(tables),
            'tables': tables,
            'metadata': metadata
        }
    
    finally:
        # Clean up temporary file
        os.unlink(tmp_file_path)

@app.post('/extract-word')
async def extract_word(file: UploadFile = File(...)):
    """Extract content from a base64-encoded Word document"""
    try:
        # Read the uploaded file
        file_content = await file.read()
        
        # If it's a plain text file with base64 content, decode it
        if file.filename.endswith('.txt'):
            base64_content = file_content.decode('utf-8')
            file_bytes = decode_base64_file(base64_content)
        else:
            # If it's already the binary file
            file_bytes = file_content
        
        # Extract content from Word document
        extracted_data = extract_word_content(file_bytes)
        
        return JSONResponse(content=extracted_data)
    
    except Exception as e:
        return JSONResponse(content={'error': str(e)}, status_code=400)

@app.post('/decode-and-extract')
async def decode_and_extract(base64_data: str):
    """Decode base64 string and extract Word document content"""
    try:
        # Decode base64 data
        file_bytes = decode_base64_file(base64_data)
        
        # Extract content from Word document
        extracted_data = extract_word_content(file_bytes)
        
        return JSONResponse(content=extracted_data)
    
    except Exception as e:
        return JSONResponse(content={'error': str(e)}, status_code=400)

if __name__ == '__main__':
    import uvicorn
    uvicorn.run(app, host='0.0.0.0', port=8001)